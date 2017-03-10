#!/usr/bin/env python3
# -*- encoding: utf-8 -*-
import docx


def test_word_read():
    doc = docx.Document('demo.docx')
    print(len(doc.paragraphs))
    print(doc.paragraphs[0].text)
    print(doc.paragraphs[1].text)
    print(len(doc.paragraphs[1].runs))
    print(doc.paragraphs[1].runs[0].text)


def test_word_full_read():
    from word_handle import readDocx
    print(readDocx.getText('demo.docx'))


def test_change_word_style():
    doc = docx.Document('demo.docx')
    print(doc.paragraphs[0].style)


if __name__ == '__main__':
    test_change_word_style()
